#!/usr/bin/env python3
"""
semi_auto_coder_v2.py — Semi-automatisches Kodieren mit Redirects, Valenz (Kontext + Heuristik)
Install:
    pip install python-docx pandas pyyaml rapidfuzz
    # optional:
    # pip install spacy && python -m spacy download de_core_news_sm

Beispiel:
    python semi_auto_coder_v2.py ./data TAP_codebook.yaml output/out.csv -r \
      -k 3 --min-score 1.0 --fuzzy-threshold 82 --w-fuzzy 1.0 --lemmatize
"""
import argparse, csv, re, sys, unicodedata
from pathlib import Path
from typing import Dict, List, Tuple, Optional

import yaml, pandas as pd
from rapidfuzz import fuzz
from docx import Document

try:
    import spacy
    _NLP = None
except Exception:
    spacy = None
    _NLP = None

# ------------ Kontext-Marker (positiv/negativ) ------------
# erkennt u.a. Spalten-/Zeilenköpfe oder Überschriften wie:
# "… positiv", "… negativ", "a positiv / b negativ", "unterstützt", "erschwert"
VAL_POS_PAT = re.compile(
    r"(^|[\s:/\\\[\]\(\)|\-])\s*(positiv(e|en|er|es)?|pos\.?|a\s*positiv|\+|unterst(ü|u)tz(t|end)?)\b",
    re.IGNORECASE
)
VAL_NEG_PAT = re.compile(
    r"(^|[\s:/\\\[\]\(\)|\-])\s*(negativ(e|en|er|es)?|neg\.?|b\s*negativ|\-|erschwert)\b",
    re.IGNORECASE
)

# ------------ Heuristik: globale Valenz-Signalwörter (ergänzbar via YAML) ------------
DEFAULT_GLOBAL_POS = [
    "hilfreich","verständlich","anschaulich","klar","strukturiert","wertschätzend","freundlich",
    "engagiert","gut vorbereitet","praxisrelevant","motiviert","angenehme Atmosphäre","übersichtlich",
    "konstruktives Feedback","angemessenes Tempo","Praxisbezug","Transferfragen","guter Überblick"
]
DEFAULT_GLOBAL_NEG = [
    "unverständlich","unklar","chaotisch","unstrukturiert","monoton","langweilig","zu schnell","zu langsam",
    "überfordert","kein Feedback","nutzloses Feedback","fehlende Korrekturen","Material fehlt",
    "Technikprobleme","ungünstige Uhrzeit","Workload unangemessen","keine Einbindung","schlecht erreichbar"
]

# Negationswörter, die die Polarität flippen können (einfaches Fenster)
NEGATIONS = r"(nicht|kein(?:e|er|en|em|es)?|ohne|mangelnd(?:e|er|es|en)?)"

Location = Tuple[str, str, Optional[str]]  # (where, text, ctx_valence)

# -------------------- Helpers --------------------
def _load_spacy():
    global _NLP
    if _NLP is None and spacy is not None:
        try:
            _NLP = spacy.load("de_core_news_sm")
        except Exception:
            _NLP = None
    return _NLP

def find_docx_paths(path: Path, recursive: bool) -> List[Path]:
    if path.is_file() and path.suffix.lower() == ".docx":
        return [path]
    if path.is_dir():
        globber = path.rglob("*") if recursive else path.glob("*")
        return [p for p in globber if p.is_file() and p.suffix.lower() == ".docx"]
    return []

def detect_ctx_valence(text: str) -> Optional[str]:
    t = (text or "").strip()
    if not t:
        return None
    if VAL_POS_PAT.search(t):
        return "positive"
    if VAL_NEG_PAT.search(t):
        return "negative"
    return None

def _iter_paragraphs_with_valence(container, prefix: str):
    current = None
    for i, p in enumerate(container.paragraphs, start=1):
        txt = p.text.strip()
        if not txt:
            continue
        v = detect_ctx_valence(txt)
        if v:
            current = v
            # reine Überschrift mit Valenzmarker wird als Kontext übernommen, aber nicht als Eintrag gewertet
            continue
        yield (f"{prefix} paragraph {i}", txt, current)

def _iter_tables_with_valence(container, prefix: str):
    for ti, table in enumerate(container.tables, start=1):
        col_val = {}
        if table.rows:
            hdr = table.rows[0]
            for c, cell in enumerate(hdr.cells, start=1):
                v = detect_ctx_valence(cell.text)
                if v:
                    col_val[c] = v
        for r, row in enumerate(table.rows, start=1):
            row_marker = detect_ctx_valence(row.cells[0].text) if row.cells else None
            for c, cell in enumerate(row.cells, start=1):
                t = cell.text.strip()
                if not t:
                    continue
                v = col_val.get(c) or row_marker or None
                yield (f"{prefix} table {ti} r{r}c{c}", t, v)

def extract_locations(doc: Document) -> List[Location]:
    out: List[Location] = []
    out.extend(_iter_paragraphs_with_valence(doc, "body"))
    out.extend(_iter_tables_with_valence(doc, "body"))
    for si, section in enumerate(doc.sections, start=1):
        header, footer = section.header, section.footer
        out.extend(_iter_paragraphs_with_valence(header, f"header s{si}"))
        out.extend(_iter_tables_with_valence(header, f"header s{si}"))
        out.extend(_iter_paragraphs_with_valence(footer, f"footer s{si}"))
        out.extend(_iter_tables_with_valence(footer, f"footer s{si}"))
    return out

def normalize(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip())

def lemmatize_de(text: str) -> str:
    nlp = _load_spacy()
    if not nlp:
        return text
    doc = nlp(text)
    return " ".join(tok.lemma_ for tok in doc)

def compile_word_regex(term: str) -> re.Pattern:
    return re.compile(rf"\b{re.escape(term)}\b", re.IGNORECASE)

# -------------------- Scoring --------------------
def apply_redirects(matches: Dict[str, float], evidence: Dict[str, List[str]], text: str, redirects: List[Dict]) -> Tuple[Dict[str,float],Dict[str,List[str]]]:
    for rd in redirects or []:
        pat, to = rd.get("pattern"), rd.get("to")
        if not pat or not to:
            continue
        try:
            if re.search(pat, text, flags=re.IGNORECASE):
                ev = evidence.get("self", [])
                ev.append(f"REDIR→{to}:{pat}")
                evidence["self"] = ev
                matches["__redirect_to__"] = to
        except re.error:
            continue
    return matches, evidence

def score_text(text: str, codebook: Dict, fuzzy_threshold: int,
               w_keyword: float, w_regex: float, w_fuzzy: float, w_excl: float,
               use_lemma: bool = False):
    raw = normalize(text)
    proc = lemmatize_de(raw) if use_lemma else raw
    lower = proc.lower()

    cat_scores: Dict[str, float] = {}
    cat_hits: Dict[str, List[str]] = {}

    for cat, rules in (codebook.get("categories") or {}).items():
        score = 0.0
        hits: List[str] = []
        redirects = rules.get("redirects", [])

        # Keywords (ganze Wörter)
        for kw in rules.get("keywords", []) or []:
            if compile_word_regex(kw).search(proc):
                score += w_keyword; hits.append(f"KW:{kw}")

        # Regex
        for rx in rules.get("regex", []) or []:
            try:
                if re.search(rx, proc, flags=re.IGNORECASE|re.MULTILINE):
                    score += w_regex; hits.append(f"RX:{rx}")
            except re.error:
                pass

        # Fuzzy
        for kw in rules.get("keywords", []) or []:
            m = fuzz.partial_ratio(lower, kw.lower())
            if m >= fuzzy_threshold:
                score += w_fuzzy; hits.append(f"FZ:{kw}({m})")

        # Excludes (ziehen ab)
        for ex in rules.get("excludes", []) or []:
            if compile_word_regex(ex).search(proc):
                score += w_excl; hits.append(f"EX:{ex}")

        # Redirects
        tmp_scores = {"self": score}
        tmp_hits = {"self": hits.copy()}
        tmp_scores, tmp_hits = apply_redirects(tmp_scores, tmp_hits, proc, redirects)
        redirect_to = tmp_scores.pop("__redirect_to__", None)

        if score != 0:
            cat_scores[cat] = score
            cat_hits[cat] = hits

        if redirect_to and score > 0:
            cat_scores[redirect_to] = max(cat_scores.get(redirect_to, 0.0), score + 0.01)
            ev = cat_hits.get(redirect_to, [])
            ev.append(f"REDIR from {cat}: " + ";".join(tmp_hits["self"]))
            cat_hits[redirect_to] = ev
            cat_scores[cat] = max(0.0, cat_scores.get(cat, 0.0) - 0.2)

        # Valenz-Signale (nur Evidenz, nicht score-wirksam)
        vs = rules.get("valence_signals", {})
        if vs:
            for tag in vs.get("positive", []):
                if re.search(re.escape(tag), proc, flags=re.IGNORECASE):
                    hits.append(f"VAL:+:{tag}")
            for tag in vs.get("negative", []):
                if re.search(re.escape(tag), proc, flags=re.IGNORECASE):
                    hits.append(f"VAL:-:{tag}")
            cat_hits[cat] = hits

    total = sum(max(s, 0) for s in cat_scores.values())
    confidences = {c: (max(s, 0)/total) if total>0 else 0.0 for c,s in cat_scores.items()}
    return cat_scores, confidences, cat_hits

# -------------------- Valenz-Inferenz --------------------
def count_valence_terms(text: str, positives: List[str], negatives: List[str]):
    """Zählt pos/neg Terme mit einfacher Negationslogik (zwei-Wort-Fenster)."""
    t = " " + unicodedata.normalize("NFKC", text) + " "
    pos_hits, neg_hits = [], []
    pos_count = neg_count = 0

    for term in positives:
        term_re = re.compile(rf"\b{re.escape(term)}\b", re.IGNORECASE)
        neg_re  = re.compile(rf"\b{NEGATIONS}\s+(?:\w+\s+){{0,2}}{re.escape(term)}\b", re.IGNORECASE)
        neg_m = list(neg_re.finditer(t))
        pos_m = [m for m in term_re.finditer(t)]
        pos_count += max(0, len(pos_m) - len(neg_m))
        neg_count += len(neg_m)
        if pos_m: pos_hits.append(term)
        if neg_m: neg_hits.append(f"nicht {term}")

    for term in negatives:
        term_re = re.compile(rf"\b{re.escape(term)}\b", re.IGNORECASE)
        neg_re  = re.compile(rf"\b{NEGATIONS}\s+(?:\w+\s+){{0,2}}{re.escape(term)}\b", re.IGNORECASE)
        neg_m = list(neg_re.finditer(t))
        neg_m_plain = [m for m in term_re.finditer(t)]
        neg_count += max(0, len(neg_m_plain) - len(neg_m))
        pos_count += len(neg_m)
        if neg_m_plain: neg_hits.append(term)
        if neg_m: pos_hits.append(f"nicht {term}")

    return pos_count, neg_count, pos_hits[:5], neg_hits[:5]

def infer_valence(ctx_valence: Optional[str], cat_hits: Dict[str, List[str]], text: str, codebook: Dict):
    # 1) HARTE KONTEXT-MARKER
    if ctx_valence in ("positive", "negative"):
        return ctx_valence, "context", f"ctx:{ctx_valence}"

    # 2) SIGNALS AUS KATEGORIEN (VAL:+:/VAL:-:)
    pos_sig = neg_sig = 0
    pos_evi = []; neg_evi = []
    for hits in cat_hits.values():
        for h in hits:
            if h.startswith("VAL:+:"):
                pos_sig += 1; pos_evi.append(h.split("VAL:+:")[-1])
            elif h.startswith("VAL:-:"):
                neg_sig += 1; neg_evi.append(h.split("VAL:-:")[-1])
    if pos_sig > neg_sig:
        return "positive", "signals", f"VAL+:{pos_sig} ({', '.join(pos_evi[:5])})"
    if neg_sig > pos_sig:
        return "negative", "signals", f"VAL-:{neg_sig} ({', '.join(neg_evi[:5])})"

    # 3) GLOBALE HEURISTIK (inkl. Negation)
    vg = (codebook.get("valence_global") or {})
    pos_list = list(vg.get("positive", [])) + DEFAULT_GLOBAL_POS
    neg_list = list(vg.get("negative", [])) + DEFAULT_GLOBAL_NEG
    p, n, ph, nh = count_valence_terms(text, pos_list, neg_list)
    if p > n:
        return "positive", "global", f"pos:{p} ({', '.join(ph)})"
    if n > p:
        return "negative", "global", f"neg:{n} ({', '.join(nh)})"

    return "neutral", "none", ""

# -------------------- Main --------------------
def main():
    ap = argparse.ArgumentParser(description="Semi-automatisches Kodieren (Redirects & Valenz)")
    ap.add_argument("inpath")
    ap.add_argument("codebook")
    ap.add_argument("outcsv")
    ap.add_argument("-r","--recursive", action="store_true")
    ap.add_argument("-k","--topk", type=int, default=3)
    ap.add_argument("--min-score", type=float, default=1.0)
    ap.add_argument("--fuzzy-threshold", type=int, default=82)
    ap.add_argument("--w-keyword", type=float, default=1.0)
    ap.add_argument("--w-regex", type=float, default=1.5)
    ap.add_argument("--w-fuzzy", type=float, default=1.0)
    ap.add_argument("--w-excl", type=float, default=-1.0)
    ap.add_argument("--lemmatize", action="store_true", help="spaCy de_core_news_sm verwenden (falls installiert)")
    args = ap.parse_args()

    inpath = Path(args.inpath).expanduser().resolve()
    codebook_path = Path(args.codebook).expanduser().resolve()
    outcsv = Path(args.outcsv).expanduser().resolve()

    with open(codebook_path, "r", encoding="utf-8") as f:
        codebook = (yaml.safe_load(f) or {"categories": {}})

    paths = find_docx_paths(inpath, args.recursive)
    if not paths:
        print("Keine .docx-Dateien gefunden.", file=sys.stderr); sys.exit(1)

    rows = []
    for doc_path in paths:
        try:
            doc = Document(str(doc_path))
        except Exception as e:
            rows.append([str(doc_path), "ERROR", "", "", 0.0, "", str(e), "", "", ""]); continue

        for where, text, ctx_val in extract_locations(doc):
            if not text.strip():
                continue

            scores, confs, hits = score_text(
                text, codebook, args.fuzzy_threshold,
                args.w_keyword, args.w_regex, args.w_fuzzy, args.w_excl,
                use_lemma=args.lemmatize
            )

            ordered = sorted(scores.items(), key=lambda x: x[1], reverse=True)
            top = [c for c, s in ordered[:args.topk] if s >= args.min_score]
            top_scores = [scores[c] for c in top]
            top_conf = [confs.get(c, 0.0) for c in top]
            top_hits = [";".join(hits.get(c, [])) for c in top]

            valence, val_source, val_evi = infer_valence(ctx_val, hits, text, codebook)

            rows.append([
                str(doc_path), where, normalize(text),
                "|".join(top),
                "|".join(f"{x:.2f}" for x in top_scores),
                "|".join(f"{x:.2f}" for x in top_conf),
                "|".join(top_hits),
                valence, val_source, val_evi
            ])

    outcsv.parent.mkdir(parents=True, exist_ok=True)
    with open(outcsv, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow([
            "file","where","text",
            "suggested_labels","scores","confidences","evidence",
            "valence","valence_source","valence_evidence"
        ])
            # ^ valence_source: context | signals | global | none
        w.writerows(rows)

    print(f"Fertig. Dateien: {len(paths)}  → {outcsv}")

if __name__ == "__main__":
    main()
